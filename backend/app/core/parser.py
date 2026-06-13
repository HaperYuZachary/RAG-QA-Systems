from dataclasses import dataclass
from pathlib import Path
import tempfile
from zipfile import ZIP_DEFLATED, ZipFile
from xml.etree import ElementTree

from app.utils.file_utils import UnsupportedFileTypeError


MIN_EXTRACTED_PDF_CHARS = 50
RELATIONSHIP_NAMESPACE = "http://schemas.openxmlformats.org/package/2006/relationships"
HYPERLINK_RELATIONSHIP_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
)


@dataclass(frozen=True)
class ParsedPage:
    page: int | None
    text: str


def parse_markdown(path: str | Path) -> list[ParsedPage]:
    text = Path(path).read_text(encoding="utf-8")
    return [ParsedPage(page=None, text=text)]


def parse_docx(path: str | Path) -> list[ParsedPage]:
    document = _open_docx_document(path)
    text_blocks = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]

    for table in document.tables:
        for row in table.rows:
            cells = [
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            ]
            if cells:
                text_blocks.append(" | ".join(cells))

    text = "\n".join(text_blocks)
    return [ParsedPage(page=None, text=text)]


def _open_docx_document(path: str | Path):
    from docx import Document

    try:
        return Document(path)
    except KeyError as exc:
        if "word/#" not in str(exc):
            raise

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            sanitized_path = Path(tmp.name)

        try:
            if not _sanitize_docx_internal_hyperlinks(path, sanitized_path):
                raise
            return Document(sanitized_path)
        finally:
            sanitized_path.unlink(missing_ok=True)


def _sanitize_docx_internal_hyperlinks(
    source_path: str | Path,
    target_path: str | Path,
) -> bool:
    removed_any = False

    with ZipFile(source_path, "r") as source_zip, ZipFile(
        target_path,
        "w",
        ZIP_DEFLATED,
    ) as target_zip:
        for item in source_zip.infolist():
            data = source_zip.read(item.filename)
            if item.filename.endswith(".rels"):
                cleaned_data, removed = _remove_invalid_bookmark_relationships(data)
                data = cleaned_data
                removed_any = removed_any or removed
            target_zip.writestr(item, data)

    return removed_any


def _remove_invalid_bookmark_relationships(data: bytes) -> tuple[bytes, bool]:
    try:
        root = ElementTree.fromstring(data)
    except ElementTree.ParseError:
        return data, False

    removed = False
    for relationship in list(root):
        if (
            relationship.get("Type") == HYPERLINK_RELATIONSHIP_TYPE
            and (relationship.get("Target") or "").startswith("#")
        ):
            root.remove(relationship)
            removed = True

    if not removed:
        return data, False

    ElementTree.register_namespace("", RELATIONSHIP_NAMESPACE)
    return ElementTree.tostring(root, encoding="utf-8", xml_declaration=True), True


def parse_pdf(path: str | Path) -> list[ParsedPage]:
    import fitz

    document = fitz.open(path)
    pages = []
    for index, page in enumerate(document):
        text = page.get_text().strip()
        if text:
            pages.append(ParsedPage(page=index + 1, text=text))
    document.close()

    total_chars = sum(len(page.text.strip()) for page in pages)
    if total_chars < MIN_EXTRACTED_PDF_CHARS:
        raise ValueError("PDF appears to be scanned or empty; OCR is not supported")

    return pages


def parse(path: str | Path, file_type: str) -> list[ParsedPage]:
    parsers = {
        "pdf": parse_pdf,
        "docx": parse_docx,
        "md": parse_markdown,
    }
    try:
        parser = parsers[file_type]
    except KeyError as exc:
        raise UnsupportedFileTypeError(f"Unsupported file type: {file_type}") from exc
    return parser(path)
